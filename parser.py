"""Document extraction and Mimo AI parsing for exam papers."""

import base64
import json
import logging
import os
import re
import tempfile
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests

from config import MIMO_API_KEY, MIMO_API_URL, MIMO_MODEL

logger = logging.getLogger(__name__)

IMAGE_MAX_DIM = 480
IMAGE_JPEG_QUALITY = 35
QUESTION_BLOCKS_FIRST_THRESHOLD = 6
QUESTION_BLOCK_WORKERS = 3

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import fitz
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _natural_sort_key(name):
    """Sort image2.png before image10.png (lexicographic puts 10 before 2)."""
    return [int(s) if s.isdigit() else s.lower()
            for s in re.split(r"(\d+)", name)]


def _compress_image_to_jpeg_bytes(data, max_dim=IMAGE_MAX_DIM, quality=IMAGE_JPEG_QUALITY):
    """Resize+recompress to JPEG to keep payloads small.

    Returns (bytes, mime). Falls back to (original_data, None) if PIL is missing
    or the source bytes can't be decoded as an image.
    """
    try:
        from PIL import Image
    except ImportError:
        return data, None

    import io
    try:
        img = Image.open(io.BytesIO(data))
        img.load()

        if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
            rgba = img.convert("RGBA")
            bg = Image.new("RGB", rgba.size, (255, 255, 255))
            bg.paste(rgba, mask=rgba.split()[-1])
            img = bg
        elif img.mode != "RGB":
            img = img.convert("RGB")

        w, h = img.size
        if max(w, h) > max_dim:
            scale = max_dim / max(w, h)
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality, optimize=True)
        return buf.getvalue(), "image/jpeg"
    except Exception:
        return data, None


IMAGE_ANNOTATION_PROMPT = """
你是高考化学试卷图片标注助手。你只输出图片标注 JSON，不重新解析整张试卷。

任务：根据已解析出的题目列表和当前批次图片，判断每张图片属于哪道题，并概括图片中的关键信息。

输出格式必须是严格 JSON：
{
  "annotations": [
    {
      "question_num": 1,
      "image_refs": [1],
      "visual_note": "图片中的关键实验装置、曲线、表格或结构信息"
    }
  ]
}

规则：
- 只输出 annotations，不输出 questions。
- image_refs 使用文档中的全局图片序号，不是当前批次内序号。
- 无法确定属于哪道题的图片不要输出。
- 不要改写题干、答案或解析。
"""


EXTRACTION_PROMPT = """你是一个高考化学试卷解析专家。请将下面的试卷内容解析为结构化题目列表。
一份标准的高考化学试卷通常包含 10-15 道题目。请只提取真正的考试题目，不要提取试卷说明、答题须知、评分标准、目录、页眉页脚等非题目内容。

对于每道题目，提取以下字段并以 JSON 数组格式返回：
- question_num: 题号，整数
- q_type: 题型，只能是 “选择题” / “填空题” / “实验题” / “计算题” / “简答题” 之一
- stem: 题干原文，完整保留，包括化学方程式、图表说明和符号；大题下所有小问合并到同一个 stem
- options: 选项列表，如 [{“A”:”...”},{“B”:”...”}]；非选择题为空数组
- answer: 标准答案。文档中已有答案或解析时，必须以文档原文为准；只有文档没有明确答案时，才根据题目内容推断最合理答案
- explanation: 解题分析。文档中已有答案或解析时，必须以文档原文为准，不要改写、替换或重新推断；只有文档没有解析时，才生成包括考查知识点、解题思路、关键步骤和易错点的解析，不少于 50 字
- topics: 涉及的知识点关键词，用空格分隔，至少 2 个
- image_refs: 该题目关联的图片序号列表，如 [1, 2, 3]。图片按在文档中出现的顺序从 1 开始编号。如果题目文字中提到”如图所示””如下图””装置图”等且附有图片，请标出对应图片序号。无图片的题目返回空列表 []

重要规则：
1. 每道题必须有 answer 和 explanation，不能留空。
2. 选择题的 answer 必须是单个选项字母。
3. 每道大题作为一个整体，不要拆成多个小题。
4. 如果 PDF 文字乱码、公式显示不正常，结合页面图片和上下文还原。
5. 如果文档后半部分、答案页、解析页、表格或图片中包含”参考答案””答案””解析””详解”等内容，先匹配到对应题号，再填入 answer 和 explanation；不得用你自己的答案覆盖文档已有内容。
6. 图片序号是全局的：如果题目提到第 3 张图片，image_refs 里填 3，不要从当前题重新开始编号。

只返回下面格式的 JSON，不要包含其他内容：
{
  “questions”: [
    {
      “question_num”: 1,
      “q_type”: “选择题”,
      “stem”: “...”,
      “options”: [{“A”:”...”},{“B”:”...”}],
      “answer”: “B”,
      “explanation”: “本题考查...”,
      “topics”: “关键字 关键字”,
      “image_refs”: [1]
    }
  ]
}
"""


def get_document_kind(filepath):
    """Return supported document kind for a path, or None if unsupported."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".doc", ".docx"}:
        return "word"
    return None


# Above this score, a candidate text is "good enough" — we stop trying the
# slower libraries (pdfplumber, PyPDF2) and return immediately. Tuned so that
# a scanned/near-empty extraction (score ~0) keeps falling back, while a real
# chemistry paper with equations/options scores well above it.
PDF_TEXT_GOOD_ENOUGH_SCORE = 50


def extract_text_from_pdf(filepath):
    """Extract text from a PDF file and keep the richest chemistry text.

    Performance: tries PyMuPDF (fitz) first and short-circuits when its output
    scores well, so the slower pdfplumber/PyPDF2 passes only run as fallback.
    """
    if HAS_FITZ:
        fitz_texts = _extract_pdf_text_with_fitz(filepath)
        best = None
        best_score = -1
        for text in fitz_texts:
            if not text or not text.strip():
                continue
            score = _score_pdf_text(text)
            if score > best_score:
                best, best_score = text, score
        if best_score >= PDF_TEXT_GOOD_ENOUGH_SCORE:
            return best

        # fitz wasn't good enough; collect it and try the other libraries.
        candidates = [t for t in fitz_texts if t and t.strip()]
    else:
        candidates = []

    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(filepath) as pdf:
                text_parts = []
                layout_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    layout_text = page.extract_text(layout=True, x_tolerance=1, y_tolerance=3)
                    if layout_text:
                        layout_parts.append(layout_text)
                candidates.append("\n\n".join(text_parts))
                candidates.append("\n\n".join(layout_parts))
        except Exception:
            pass

    if HAS_PYPDF2:
        try:
            text_parts = []
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            candidates.append("\n\n".join(text_parts))
        except Exception:
            pass

    return choose_best_pdf_text(candidates)


def _extract_pdf_text_with_fitz(filepath):
    """Return [plain_text, block_text] from PyMuPDF, or [] on failure."""
    try:
        doc = fitz.open(filepath)
        try:
            text_parts = []
            block_parts = []
            for page in doc:
                text = page.get_text("text")
                if text:
                    text_parts.append(text)
                blocks = page.get_text("blocks")
                if blocks:
                    blocks = sorted(blocks, key=lambda b: (round(b[1]), b[0]))
                    block_parts.extend(b[4] for b in blocks if len(b) > 4 and b[4].strip())
            return ["\n\n".join(text_parts), "\n".join(block_parts)]
        finally:
            doc.close()
    except Exception:
        return []


def choose_best_pdf_text(candidates):
    """Choose the extracted PDF text that best preserves chemistry questions."""
    clean_candidates = [c for c in candidates if c and c.strip()]
    if not clean_candidates:
        return None
    return max(clean_candidates, key=_score_pdf_text)


def _score_pdf_text(text):
    compact = re.sub(r"\s+", "", text)
    score = len(text) * 0.01
    score += len(re.findall(r"[A-D][\.．、]", text)) * 30
    score += len(re.findall(r"[①②③④⑤⑥⑦⑧⑨⑩]", text)) * 20
    score += len(re.findall(r"ΔH|kJ/mol|mol|aq|g\)|l\)|s\)", text)) * 18
    score += len(re.findall(r"[A-Z][a-z]?\d*", compact)) * 3
    score += len(re.findall(r"[+=→⇌]", text)) * 5
    score += len(re.findall(r"N2H4|NH3|H2O|CO2|H2SO4|NaOH", compact)) * 25
    return score


def render_pdf_pages_as_data_urls(filepath, max_pages=8, zoom=1.8):
    """Render PDF pages to PNG data URLs for multimodal model parsing."""
    if not HAS_FITZ:
        return []

    images = []
    try:
        doc = fitz.open(filepath)
        try:
            for page in doc[:max_pages]:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                encoded = base64.b64encode(pixmap.tobytes("png")).decode("ascii")
                images.append(f"data:image/png;base64,{encoded}")
        finally:
            doc.close()
    except Exception:
        return []
    return images


def extract_pdf_embedded_images_as_data_urls(filepath, max_images=None):
    """Extract images embedded inside a PDF as data URLs.

    Each image is recompressed to JPEG (max 1280px, quality 75) so request
    bodies stay small even when the document has many high-res figures.

    All embedded images are returned by default (max_images=None); pass an
    int only to cap the count.
    """
    if not HAS_FITZ:
        return []

    images = []
    seen = set()
    try:
        doc = fitz.open(filepath)
        try:
            positioned_images = []
            for page_index in range(len(doc)):
                page = doc[page_index]
                for image_info in page.get_images(full=True):
                    xref = image_info[0]
                    if xref in seen:
                        continue
                    seen.add(xref)
                    rects = page.get_image_rects(xref)
                    if rects:
                        rect = rects[0]
                        sort_key = (page_index, rect.y0, rect.x0)
                    else:
                        sort_key = (page_index, float("inf"), float("inf"))
                    positioned_images.append((sort_key, xref))
            for _sort_key, xref in sorted(positioned_images):
                image = doc.extract_image(xref)
                image_bytes = image.get("image")
                if not image_bytes:
                    continue
                compressed, jpeg_mime = _compress_image_to_jpeg_bytes(image_bytes)
                if jpeg_mime:
                    final_bytes, final_mime = compressed, jpeg_mime
                else:
                    ext = image.get("ext", "png").lower()
                    if ext in {"jpg", "jpeg"}:
                        final_mime = "image/jpeg"
                    elif ext == "webp":
                        final_mime = "image/webp"
                    else:
                        final_mime = "image/png"
                    final_bytes = image_bytes
                encoded = base64.b64encode(final_bytes).decode("ascii")
                images.append(f"data:{final_mime};base64,{encoded}")
                if max_images is not None and len(images) >= max_images:
                    return images
        finally:
            doc.close()
    except Exception:
        return []
    return images


def extract_docx_embedded_images_as_data_urls(filepath, max_images=None):
    """Extract images embedded inside a .docx (OOXML zip) as data URLs.

    Sorts by natural numeric order so image2.png precedes image10.png. Each
    image is recompressed to JPEG so the model receives smaller payloads.

    All embedded images are returned by default (max_images=None); pass an
    int only to cap the count.
    """
    images = []
    mime_map = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }
    try:
        with zipfile.ZipFile(filepath) as zf:
            media_names = sorted(
                (name for name in zf.namelist()
                 if name.startswith("word/media/")
                 and os.path.splitext(name)[1].lower() in mime_map),
                key=_natural_sort_key,
            )
            try:
                rels = zf.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
                document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
                rel_targets = dict(re.findall(r'<Relationship[^>]+Id="([^"]+)"[^>]+Target="media/([^"]+)"', rels))
                ordered = []
                for rel_id in re.findall(r'r:embed="([^"]+)"', document_xml):
                    target = rel_targets.get(rel_id)
                    if target:
                        name = "word/media/" + target
                        if name in media_names and name not in ordered:
                            ordered.append(name)
                media_names = ordered + [name for name in media_names if name not in ordered]
            except Exception:
                pass
            for name in media_names:
                ext = os.path.splitext(name)[1].lower()
                with zf.open(name) as fh:
                    data = fh.read()
                if not data:
                    continue
                compressed, jpeg_mime = _compress_image_to_jpeg_bytes(data)
                if jpeg_mime:
                    final_bytes, final_mime = compressed, jpeg_mime
                else:
                    final_bytes, final_mime = data, mime_map[ext]
                encoded = base64.b64encode(final_bytes).decode("ascii")
                images.append(f"data:{final_mime};base64,{encoded}")
                if max_images is not None and len(images) >= max_images:
                    break
    except (zipfile.BadZipFile, KeyError, OSError):
        return []
    return images


def extract_text_from_word(filepath):
    """Extract text from docx files, with a best-effort Windows fallback for doc."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        if not HAS_DOCX:
            return None
        try:
            doc = Document(filepath)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append("\t".join(cells))
            return "\n".join(parts) or None
        except Exception:
            return None

    if ext == ".doc":
        return _extract_text_from_legacy_doc(filepath)

    return None


def _extract_text_from_legacy_doc(filepath):
    """Use Microsoft Word COM automation when available to read legacy .doc."""
    try:
        import win32com.client
    except ImportError:
        return None

    word = None
    temp_path = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
        fd, temp_path = tempfile.mkstemp(suffix=".txt")
        os.close(fd)
        doc.SaveAs(temp_path, FileFormat=2)
        doc.Close(False)
        with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip() or None
    except Exception:
        return None
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except OSError:
                pass


def extract_text_from_document(filepath):
    """Extract text from a supported document path."""
    kind = get_document_kind(filepath)
    if kind == "pdf":
        return extract_text_from_pdf(filepath)
    if kind == "word":
        return extract_text_from_word(filepath)
    return None


MIMO_TIMEOUT_SECONDS = 180


def call_mimo(paper_text, image_urls=None, system_prompt=EXTRACTION_PROMPT, retry=True):
    """Send extraction request to the Mimo OpenAI-compatible API."""
    if not MIMO_API_KEY:
        return {"error": "缺少 MIMO_API_KEY 环境变量，无法调用 Mimo 模型"}

    user_content = paper_text[:40000]
    if image_urls:
        user_content = [{"type": "text", "text": paper_text[:20000]}]
        user_content.extend(
            {"type": "image_url", "image_url": {"url": image_url}}
            for image_url in image_urls
        )

    payload = {
        "model": MIMO_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        "response_format": {"type": "json_object"},
        "max_tokens": 16000,
        "temperature": 0.1,
    }
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MIMO_API_KEY}",
    }
    try:
        resp = requests.post(MIMO_API_URL, headers=headers, json=payload, timeout=MIMO_TIMEOUT_SECONDS)
        if resp.status_code != 200:
            return {"error": f"API HTTP {resp.status_code}", "raw": resp.text[:300]}
        result = resp.json()
        content = result["choices"][0]["message"]["content"]
        return _parse_json_content(content)
    except json.JSONDecodeError:
        if retry:
            return call_mimo(paper_text[:20000], image_urls=None, system_prompt=system_prompt, retry=False)
        return {"error": "JSON parse failed"}
    except requests.Timeout as e:
        if retry:
            return call_mimo(paper_text[:20000], image_urls=None, system_prompt=system_prompt, retry=False)
        return {"error": f"API call timed out after {MIMO_TIMEOUT_SECONDS}s: {e}"}
    except (KeyError, requests.RequestException) as e:
        return {"error": f"API call failed: {e}"}


def _parse_json_content(content):
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", content)
    if m:
        content = m.group(1)
    else:
        m = re.search(r"\{[\s\S]*\}", content)
        if m:
            content = m.group()
    return json.loads(content)


def chunk_items(items, chunk_size):
    """Split a list into fixed-size chunks, keeping order and dropping nothing."""
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    return [items[i:i + chunk_size] for i in range(0, len(items), chunk_size)]


def normalize_parse_result(result, image_count=0, allowed_question_nums=None):
    """Drop unusable questions and sanitize fields the review UI depends on."""
    if not result or "questions" not in result:
        return result

    questions = []
    for question in result.get("questions", []):
        question_num = question.get("question_num")
        if not isinstance(question_num, int):
            continue
        if not isinstance(question.get("options", []), list):
            question["options"] = []
        image_refs = []
        for ref in question.get("image_refs", []):
            if type(ref) is int and 1 <= ref <= image_count and ref not in image_refs:
                image_refs.append(ref)
        if image_refs:
            question["image_refs"] = image_refs
        else:
            question.pop("image_refs", None)
        questions.append(question)
    result["questions"] = questions
    return result

def merge_question_batches(batch_results):
    """Merge per-batch parse results, deduplicating by question_num.

    Later batches override earlier ones for the same question_num: image-aware
    batches run after the text-only batch, so their versions are more complete.
    """
    by_num = {}
    order = []
    for result in batch_results:
        for question in result.get("questions", []):
            question_num = question.get("question_num")
            if question_num is None:
                continue
            if question_num not in by_num:
                order.append(question_num)
            by_num[question_num] = question
    return {"questions": [by_num[num] for num in order]}


def split_text_into_question_blocks(text):
    """Split extracted paper text into coarse question blocks by leading numbers."""
    matches = list(re.finditer(r"(?m)^\s*(\d{1,2})[\.．、]\s*", text))
    if len(matches) < 2:
        return []

    blocks = []
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        block = text[start:end].strip()
        if len(block) >= 20:
            blocks.append(block)
    return blocks


def parse_question_blocks_with_mimo(text):
    """Parse question blocks independently and merge successful results."""
    blocks = split_text_into_question_blocks(text)
    if not blocks:
        return None

    def _process_block(index, block):
        logger.info("mimo_question_block start index=%s total=%s chars=%s", index, len(blocks), len(block))
        block_start = time.perf_counter()
        result = call_mimo(block, image_urls=None, retry=False)
        if result and "error" in result:
            result = call_mimo(block, image_urls=None, retry=False)
        if result and "questions" in result:
            result = {**result, "questions": [dict(question) for question in result.get("questions", [])]}
            match = re.match(r"\s*(\d{1,2})", block)
            if match and len(result.get("questions", [])) == 1:
                result["questions"][0]["question_num"] = int(match.group(1))
        logger.info(
            "mimo_question_block done index=%s total=%s chars=%s elapsed=%.2fs error=%s",
            index,
            len(blocks),
            len(block),
            time.perf_counter() - block_start,
            bool(result and "error" in result),
        )
        return result

    results_by_index = {}
    with ThreadPoolExecutor(max_workers=QUESTION_BLOCK_WORKERS) as executor:
        futures = {
            executor.submit(_process_block, index, block): index
            for index, block in enumerate(blocks, start=1)
        }
        for future in as_completed(futures):
            index = futures[future]
            try:
                results_by_index[index] = future.result()
            except Exception as e:
                results_by_index[index] = {"error": f"Thread failed: {e}"}

    batch_results = []
    for index in range(1, len(blocks) + 1):
        result = results_by_index.get(index)
        if result and "questions" in result:
            batch_results.append(result)

    if not batch_results:
        return None
    block_question_nums = {int(m.group(1)) for block in blocks for m in [re.match(r"\s*(\d{1,2})", block)] if m}
    return normalize_parse_result(merge_question_batches(batch_results), allowed_question_nums=block_question_nums or None)


def merge_image_annotations(text_result, annotation_results, image_count):
    """Merge image-only annotations into text-parsed questions."""
    if not text_result or "questions" not in text_result:
        return text_result

    for question in text_result.get("questions", []):
        image_refs = []
        for ref in question.get("image_refs", []):
            if type(ref) is int and 1 <= ref <= image_count and ref not in image_refs:
                image_refs.append(ref)
        if image_refs:
            question["image_refs"] = image_refs
        else:
            question.pop("image_refs", None)

    questions_by_num = {
        question.get("question_num"): question
        for question in text_result.get("questions", [])
        if question.get("question_num") is not None
    }

    for result in annotation_results:
        if not result or "error" in result:
            continue
        for annotation in result.get("annotations", []):
            question = questions_by_num.get(annotation.get("question_num"))
            if question is None:
                continue

            image_refs = []
            for ref in question.get("image_refs", []):
                if type(ref) is int and 1 <= ref <= image_count and ref not in image_refs:
                    image_refs.append(ref)
            for ref in annotation.get("image_refs", []):
                if type(ref) is int and 1 <= ref <= image_count and ref not in image_refs:
                    image_refs.append(ref)
            if image_refs:
                question["image_refs"] = image_refs
            else:
                question.pop("image_refs", None)

            visual_note = annotation.get("visual_note")
            if visual_note:
                image_note = f"图像信息：{visual_note}"
                explanation = question.get("explanation")
                question["explanation"] = f"{explanation}\n\n{image_note}" if explanation else image_note

    return text_result


# Per-batch image count. Keep single multimodal calls small enough to avoid
# common provider read timeouts on image-heavy Word documents.
IMAGE_BATCH_SIZE = 4
MAX_ANNOTATION_IMAGES = 8


def call_mimo_image_annotation_batches(paper_text, text_result, image_urls, batch_size=IMAGE_BATCH_SIZE):
    """Send image-only annotation requests in chunks without reparsing questions."""
    if not image_urls:
        return []

    question_lines = []
    for question in text_result.get("questions", []):
        question_num = question.get("question_num")
        stem = str(question.get("stem", "")).replace("\r", " ").replace("\n", " ")[:120]
        question_lines.append(f"{question_num}. {stem}")

    annotation_results = []
    batches = chunk_items(list(enumerate(image_urls[:MAX_ANNOTATION_IMAGES], start=1)), batch_size)

    def _process_batch(index, batch):
        batch_refs = [ref for ref, _ in batch]
        batch_urls = [url for _, url in batch]
        prompt = (
            "只输出图片标注 JSON，不要重新解析试卷。\n\n"
            "已解析题目列表：\n"
            f"{chr(10).join(question_lines)}\n\n"
            f"当前批次全局图片序号：{batch_refs}\n\n"
            "优先匹配相邻题号：图片通常属于它前后最近的题目，不要跨很远题号，除非图片文字明确写了题号。\n"
            "输出 schema：annotations 数组，每项包含 question_num、image_refs、visual_note。\n"
            "image_refs 必须使用文档中的全局图片序号。"
        )
        logger.info("mimo_image_annotation start index=%s total=%s images=%s refs=%s", index, len(batches), len(batch_urls), batch_refs)
        batch_start = time.perf_counter()
        result = call_mimo(prompt, image_urls=batch_urls, system_prompt=IMAGE_ANNOTATION_PROMPT, retry=False)
        logger.info(
            "mimo_image_annotation done index=%s total=%s images=%s elapsed=%.2fs error=%s",
            index,
            len(batches),
            len(batch_urls),
            time.perf_counter() - batch_start,
            bool(result and "error" in result),
        )
        return result

    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = {executor.submit(_process_batch, index, batch): index for index, batch in enumerate(batches, start=1)}
        results_by_index = {}
        for future in as_completed(futures):
            idx = futures[future]
            try:
                results_by_index[idx] = future.result()
            except Exception as e:
                results_by_index[idx] = {"error": f"Thread failed: {e}"}
        
        for i in range(1, len(batches) + 1):
            annotation_results.append(results_by_index.get(i))

    return annotation_results


def call_mimo_with_image_batches(paper_text, image_urls, batch_size=IMAGE_BATCH_SIZE):
    """Send images to the model in chunks, then merge the question outputs.

    Each chunk is a separate call_mimo() round trip with only its slice of
    images attached, so no single request carries too many high-res figures
    (which is what made the old single-call path slow and prone to timeouts).
    Image ordering is the global document order, so image_refs returned by the
    model map back to the right embedded figure.
    """
    if not image_urls:
        return call_mimo(paper_text)

    batch_results = []
    batches = chunk_items(image_urls, batch_size)

    def _process_batch(index, batch):
        batch_text = (
            f"{paper_text}\n\n"
            f"Attached images batch {index}/{len(batches)}. "
            "Parse all questions supported by this batch. If a question appears "
            "in multiple batches, return the most complete version using the "
            "visible diagrams, tables, equations, and labels."
        )
        logger.info("mimo_batch start index=%s total=%s images=%s", index, len(batches), len(batch))
        batch_start = time.perf_counter()
        result = call_mimo(batch_text, image_urls=batch)
        logger.info(
            "mimo_batch done index=%s total=%s images=%s elapsed=%.2fs error=%s",
            index,
            len(batches),
            len(batch),
            time.perf_counter() - batch_start,
            bool(result and "error" in result),
        )
        return result

    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = {executor.submit(_process_batch, index, batch): index for index, batch in enumerate(batches, start=1)}
        results_by_index = {}
        for future in as_completed(futures):
            idx = futures[future]
            try:
                results_by_index[idx] = future.result()
            except Exception as e:
                results_by_index[idx] = {"error": f"Thread failed: {e}"}
        
        for i in range(1, len(batches) + 1):
            res = results_by_index.get(i)
            if res and "error" in res:
                return res
            if res:
                batch_results.append(res)

    return normalize_parse_result(merge_question_batches(batch_results))


def call_deepseek(paper_text, retry=True):
    """Compatibility wrapper for older imports."""
    return call_mimo(paper_text, retry=retry)


def save_images_to_disk(image_data_urls, paper_id, images_dir):
    """Decode base64 data URLs and save them as JPEG files on disk.

    Returns a list of ``(seq, file_path)`` tuples where *seq* is the 1-based
    sequence number that matches the order the images appeared in the document
    (and the ``image_refs`` values the AI model was told to use).
    """
    if not image_data_urls:
        return []

    os.makedirs(images_dir, exist_ok=True)
    saved = []
    for seq, data_url in enumerate(image_data_urls, start=1):
        try:
            # Parse "data:<mime>;base64,<payload>"
            header, payload = data_url.split(",", 1)
            raw_bytes = base64.b64decode(payload)

            # Always save as JPEG — the data is already JPEG-compressed by
            # _compress_image_to_jpeg_bytes, so just write it through.
            filename = f"{paper_id}_{seq}.jpg"
            filepath = os.path.join(images_dir, filename)
            with open(filepath, "wb") as f:
                f.write(raw_bytes)
            saved.append((seq, filepath))
        except Exception:
            continue
    return saved


def parse_document_to_questions(filepath):
    """Parse a PDF, DOC, or DOCX file into structured question data."""
    total_start = time.perf_counter()
    kind = get_document_kind(filepath)
    logger.info("parse_document start file=%s kind=%s", os.path.basename(filepath), kind)
    if not kind:
        return {"error": "仅支持 PDF、DOC、DOCX 文档"}

    text_start = time.perf_counter()
    text = extract_text_from_document(filepath) or ""
    logger.info(
        "parse_document text_extracted file=%s kind=%s text_chars=%s elapsed=%.2fs",
        os.path.basename(filepath),
        kind,
        len(text),
        time.perf_counter() - text_start,
    )
    image_urls = []
    image_source = None

    if kind == "pdf":
        image_urls = extract_pdf_embedded_images_as_data_urls(filepath)
        if image_urls:
            image_source = "PDF embedded images"
        if not text.strip() and not image_urls:
            image_urls = render_pdf_pages_as_data_urls(filepath)
            if image_urls:
                image_source = "PDF page images"
    elif kind == "word" and filepath.lower().endswith(".docx"):
        image_urls = extract_docx_embedded_images_as_data_urls(filepath)
        if image_urls:
            image_source = "DOCX embedded images"

    batches = len(chunk_items(image_urls, IMAGE_BATCH_SIZE)) if image_urls else 0
    logger.info(
        "parse_document images_extracted file=%s kind=%s images=%s image_source=%s batches=%s",
        os.path.basename(filepath),
        kind,
        len(image_urls),
        image_source,
        batches,
    )

    if not text.strip() and not image_urls:
        return {"error": "文档文字提取失败，文件可能为扫描件、加密文件或旧版 Word 文档缺少转换组件"}

    if len(text.strip()) < 50 and not image_urls:
        return {"error": "文档文字内容过少"}

    if not text.strip() and image_urls:
        prompt_text = f"{image_source} are attached. " + "Use them to recover diagrams, equations, tables, and visual question content."
        ai_start = time.perf_counter()
        result = call_mimo_with_image_batches(prompt_text, image_urls)
        logger.info(
            "parse_document image_only_ai_done file=%s kind=%s images=%s batches=%s elapsed=%.2fs error=%s",
            os.path.basename(filepath),
            kind,
            len(image_urls),
            batches,
            time.perf_counter() - ai_start,
            bool(result and "error" in result),
        )
        if not result:
            return {"error": "AI 解析失败，请重试"}
        if "error" in result:
            return result
        questions = result.get("questions", [])
        logger.info(
            "parse_document done file=%s kind=%s text_chars=%s images=%s batches=%s questions=%s elapsed=%.2fs",
            os.path.basename(filepath),
            kind,
            len(text),
            len(image_urls),
            batches,
            len(questions),
            time.perf_counter() - total_start,
        )
        return {
            "questions": questions,
            "raw_text": text[:5000],
            "images": image_urls,
            "image_source": image_source,
        }

    text_blocks = split_text_into_question_blocks(text)
    result = None
    if len(text_blocks) >= QUESTION_BLOCKS_FIRST_THRESHOLD:
        block_start = time.perf_counter()
        result = parse_question_blocks_with_mimo(text)
        logger.info(
            "parse_document block_ai_done file=%s kind=%s blocks=%s elapsed=%.2fs error=%s",
            os.path.basename(filepath),
            kind,
            len(text_blocks),
            time.perf_counter() - block_start,
            not bool(result),
        )

    ai_start = time.perf_counter()
    if result is None:
        result = call_mimo(text)
    logger.info(
        "parse_document text_ai_done file=%s kind=%s images=%s batches=%s elapsed=%.2fs error=%s",
        os.path.basename(filepath),
        kind,
        len(image_urls),
        batches,
        time.perf_counter() - ai_start,
        bool(result and "error" in result),
    )
    if not result:
        return {"error": "AI 解析失败，请重试"}
    if "error" in result:
        fallback_result = parse_question_blocks_with_mimo(text)
        if fallback_result:
            result = fallback_result
        else:
            return result

    if image_urls:
        image_ai_start = time.perf_counter()
        annotation_results = call_mimo_image_annotation_batches(text, result, image_urls)
        annotation_errors = sum(1 for item in annotation_results if item and "error" in item)
        result = merge_image_annotations(result, annotation_results, len(image_urls))
        logger.info(
            "parse_document image_ai_done file=%s kind=%s images=%s batches=%s annotation_errors=%s elapsed=%.2fs",
            os.path.basename(filepath),
            kind,
            len(image_urls),
            batches,
            annotation_errors,
            time.perf_counter() - image_ai_start,
        )

    result = normalize_parse_result(result, len(image_urls))
    questions = result.get("questions", [])
    logger.info(
        "parse_document done file=%s kind=%s text_chars=%s images=%s batches=%s questions=%s elapsed=%.2fs",
        os.path.basename(filepath),
        kind,
        len(text),
        len(image_urls),
        batches,
        len(questions),
        time.perf_counter() - total_start,
    )
    return {
        "questions": questions,
        "raw_text": text[:5000],
        "images": image_urls,
        "image_source": image_source,
    }


def parse_pdf_to_questions(filepath):
    """Compatibility wrapper for older PDF-only callers."""
    return parse_document_to_questions(filepath)
