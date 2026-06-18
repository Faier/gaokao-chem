import os
import tempfile
import unittest
from unittest import mock

import parser


class DocumentParserTests(unittest.TestCase):
    def test_image_batch_size_is_small_enough_to_avoid_api_timeouts(self):
        self.assertEqual(parser.IMAGE_BATCH_SIZE, 4)

    def test_image_compression_defaults_are_low_bandwidth(self):
        self.assertEqual(parser.IMAGE_MAX_DIM, 640)
        self.assertEqual(parser.IMAGE_JPEG_QUALITY, 50)

    def test_extraction_prompt_prefers_document_answers_and_explanations(self):
        self.assertIn("文档中已有答案或解析时", parser.EXTRACTION_PROMPT)
        self.assertIn("必须以文档原文为准", parser.EXTRACTION_PROMPT)
        self.assertIn("不要改写、替换或重新推断", parser.EXTRACTION_PROMPT)

    def test_choose_best_pdf_text_prefers_chemistry_rich_extraction(self):
        poor = "是一种常见燃料，可以用于火箭助推器。已知：\n①\n②\n③\n则 为\nA."
        rich = (
            "N2H4是一种常见燃料，可以用于火箭助推器。已知：\n"
            "①2NH3(g)+3N2O(g)=4N2(g)+3H2O(l) ΔH1=akJ/mol\n"
            "②N2O(g)+3H2(g)=N2H4(l)+H2O(l) ΔH2=bkJ/mol\n"
            "③N2H4(l)+9H2(g)+4O2(g)=2NH3(g)+8H2O(l) ΔH3=ckJ/mol\n"
            "则N2H4(l)+O2(g)=N2(g)+2H2O(l) ΔH为\n"
            "A. 1/4(a+3b-c)\nB. 1/4(a-3b+c)\nC. (a-3b+c)\nD. (a+3b-c)"
        )

        self.assertEqual(parser.choose_best_pdf_text([poor, rich]), rich)

    def test_document_extension_detection_supports_pdf_doc_and_docx(self):
        self.assertEqual(parser.get_document_kind("paper.PDF"), "pdf")
        self.assertEqual(parser.get_document_kind("paper.doc"), "word")
        self.assertEqual(parser.get_document_kind("paper.docx"), "word")
        self.assertIsNone(parser.get_document_kind("paper.txt"))

    def test_parse_document_uses_text_first_then_pdf_image_annotations(self):
        text_result = {"questions": [{"question_num": 1, "stem": "text question"}]}
        annotated = {"questions": [{"question_num": 1, "stem": "text question", "image_refs": [1]}]}

        with mock.patch.object(parser, "extract_text_from_pdf", return_value="题干文字" * 20), \
             mock.patch.object(parser, "extract_pdf_embedded_images_as_data_urls", return_value=["data:image/png;base64,embedded"]), \
             mock.patch.object(parser, "render_pdf_pages_as_data_urls") as render_pages, \
             mock.patch.object(parser, "call_mimo", return_value=text_result) as call_mimo, \
             mock.patch.object(parser, "call_mimo_image_annotation_batches", return_value=[{"annotations": []}]) as annotate, \
             mock.patch.object(parser, "merge_image_annotations", return_value=annotated) as merge:
            result = parser.parse_document_to_questions("paper.pdf")

        self.assertEqual(result["questions"], annotated["questions"])
        render_pages.assert_not_called()
        self.assertEqual(call_mimo.call_count, 1)
        self.assertEqual(call_mimo.call_args.args[0], "题干文字" * 20)
        annotate.assert_called_once_with("题干文字" * 20, text_result, ["data:image/png;base64,embedded"])
        merge.assert_called_once_with(text_result, [{"annotations": []}], 1)

    def test_parse_document_falls_back_to_question_blocks_when_full_text_times_out(self):
        text = "1. first question\nA. a\nB. b\n2. second question\nA. a\nB. b"
        timeout_result = {"error": "API call timed out after 180s: slow"}
        first_result = {"questions": [{"question_num": 1, "stem": "first"}]}
        second_result = {"questions": [{"question_num": 2, "stem": "second"}]}

        with mock.patch.object(parser, "extract_text_from_word", return_value=text), \
             mock.patch.object(parser, "extract_docx_embedded_images_as_data_urls", return_value=[]), \
             mock.patch.object(
                 parser,
                 "call_mimo",
                 side_effect=[timeout_result, first_result, second_result],
             ) as call_mimo:
            result = parser.parse_document_to_questions("paper.docx")

        self.assertEqual([q["question_num"] for q in result["questions"]], [1, 2])
        self.assertEqual(call_mimo.call_count, 3)
        self.assertIn("1. first question", call_mimo.call_args_list[1].args[0])
        self.assertIn("2. second question", call_mimo.call_args_list[2].args[0])

    def test_parse_document_uses_question_blocks_first_for_many_blocks(self):
        text = "\n".join(f"{i}. question {i}\nA. a\nB. b" for i in range(1, 7))
        block_result = {"questions": [{"question_num": 1, "stem": "first"}]}

        with mock.patch.object(parser, "extract_text_from_word", return_value=text), \
             mock.patch.object(parser, "extract_docx_embedded_images_as_data_urls", return_value=[]), \
             mock.patch.object(parser, "parse_question_blocks_with_mimo", return_value=block_result) as parse_blocks, \
             mock.patch.object(parser, "call_mimo") as call_mimo:
            result = parser.parse_document_to_questions("paper.docx")

        self.assertEqual(result["questions"], block_result["questions"])
        parse_blocks.assert_called_once_with(text)
        call_mimo.assert_not_called()

    def test_question_blocks_are_submitted_with_configured_concurrency(self):
        text = "\n".join(f"{i}. question {i}\nA. a\nB. b" for i in range(1, 7))
        submitted = []

        class ImmediateFuture:
            def __init__(self, result):
                self._result = result

            def result(self):
                return self._result

        class RecordingExecutor:
            def __init__(self, max_workers):
                self.max_workers = max_workers

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb):
                return False

            def submit(self, fn, index, block):
                submitted.append((self.max_workers, index, block))
                return ImmediateFuture(fn(index, block))

        with mock.patch.object(parser, "ThreadPoolExecutor", RecordingExecutor), \
             mock.patch.object(parser, "as_completed", side_effect=lambda futures: futures), \
             mock.patch.object(parser, "call_mimo", return_value={"questions": [{"question_num": 1}]}):
            result = parser.parse_question_blocks_with_mimo(text)

        self.assertEqual(result["questions"][0]["question_num"], 1)
        self.assertEqual({item[0] for item in submitted}, {parser.QUESTION_BLOCK_WORKERS})
        self.assertEqual(len(submitted), 6)

    def test_parse_document_renders_pdf_pages_and_uses_legacy_image_batching_when_text_is_missing(self):
        expected = {"questions": [{"question_num": 1, "stem": "image question"}]}

        with mock.patch.object(parser, "extract_text_from_pdf", return_value=""), \
             mock.patch.object(parser, "extract_pdf_embedded_images_as_data_urls", return_value=[]), \
             mock.patch.object(parser, "render_pdf_pages_as_data_urls", return_value=["data:image/png;base64,abc"]), \
             mock.patch.object(parser, "call_mimo_with_image_batches", return_value=expected) as batched, \
             mock.patch.object(parser, "call_mimo_image_annotation_batches") as annotate:
            result = parser.parse_document_to_questions("scan.pdf")

        self.assertEqual(result["questions"], expected["questions"])
        annotate.assert_not_called()
        args, kwargs = batched.call_args
        self.assertIn("PDF page images", args[0])
        self.assertEqual(args[1], ["data:image/png;base64,abc"])

    def test_call_mimo_sends_openai_compatible_multimodal_payload(self):
        response = mock.Mock()
        response.status_code = 200
        response.json.return_value = {
            "choices": [
                {"message": {"content": '{"questions":[{"question_num":1}]}' }}
            ]
        }

        with mock.patch.object(parser, "MIMO_API_KEY", "test-key"), \
             mock.patch.object(parser.requests, "post", return_value=response) as post:
            result = parser.call_mimo("question text", image_urls=["data:image/png;base64,abc"])

        self.assertEqual(result["questions"][0]["question_num"], 1)
        payload = post.call_args.kwargs["json"]
        self.assertEqual(payload["model"], parser.MIMO_MODEL)
        content = payload["messages"][1]["content"]
        self.assertEqual(content[0]["type"], "text")
        self.assertEqual(content[1]["type"], "image_url")

    def test_call_mimo_uses_configured_timeout_constant(self):
        response = mock.Mock()
        response.status_code = 200
        response.json.return_value = {
            "choices": [
                {"message": {"content": '{"questions":[{"question_num":1}]}' }}
            ]
        }

        with mock.patch.object(parser, "MIMO_API_KEY", "test-key"), \
             mock.patch.object(parser.requests, "post", return_value=response) as post:
            parser.call_mimo("question text")

        self.assertEqual(parser.MIMO_TIMEOUT_SECONDS, 180)
        self.assertEqual(post.call_args.kwargs["timeout"], parser.MIMO_TIMEOUT_SECONDS)

    def test_call_mimo_retries_timeout_with_smaller_text_only_payload(self):
        response = mock.Mock()
        response.status_code = 200
        response.json.return_value = {
            "choices": [
                {"message": {"content": '{"questions":[{"question_num":1}]}' }}
            ]
        }
        long_text = "question text " * 4000

        with mock.patch.object(parser, "MIMO_API_KEY", "test-key"), \
             mock.patch.object(
                 parser.requests,
                 "post",
                 side_effect=[parser.requests.ReadTimeout("slow"), response],
             ) as post:
            result = parser.call_mimo(long_text, image_urls=["data:image/png;base64,abc"])

        self.assertEqual(result["questions"][0]["question_num"], 1)
        self.assertEqual(post.call_count, 2)
        retry_payload = post.call_args_list[1].kwargs["json"]
        retry_content = retry_payload["messages"][1]["content"]
        self.assertIsInstance(retry_content, str)
        self.assertEqual(retry_content, long_text[:20000])

    def test_extract_text_from_docx_reads_paragraphs_and_tables(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "paper.docx")
            doc = Document()
            doc.add_paragraph("第一题")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "正确"
            doc.save(path)

            text = parser.extract_text_from_word(path)

        self.assertIn("第一题", text)
        self.assertIn("A", text)
        self.assertIn("正确", text)

    def test_extract_docx_embedded_images_returns_data_urls(self):
        import io
        import zipfile

        try:
            from PIL import Image
        except ImportError:
            self.skipTest("Pillow is not installed")

        def real_png(color):
            buf = io.BytesIO()
            Image.new("RGB", (40, 40), color).save(buf, format="PNG")
            return buf.getvalue()

        def real_jpeg(color):
            buf = io.BytesIO()
            Image.new("RGB", (40, 40), color).save(buf, format="JPEG")
            return buf.getvalue()

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "paper.docx")
            with zipfile.ZipFile(path, "w") as zf:
                zf.writestr("word/document.xml", "<doc/>")
                zf.writestr("word/media/image1.png", real_png((255, 0, 0)))
                zf.writestr("word/media/image2.jpeg", real_jpeg((0, 255, 0)))
                zf.writestr("word/media/notes.txt", b"ignore me")

            urls = parser.extract_docx_embedded_images_as_data_urls(path)

        self.assertEqual(len(urls), 2)
        # Both compressed to JPEG by default
        self.assertTrue(all(u.startswith("data:image/jpeg;base64,") for u in urls))

    def test_extract_docx_embedded_images_uses_natural_sort(self):
        """image2.png must come before image10.png even though lex sort puts 10 first."""
        import zipfile

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "paper.docx")
            with zipfile.ZipFile(path, "w") as zf:
                zf.writestr("word/document.xml", "<doc/>")
                # Names chosen so lex order would put image10 before image2.
                # Use raw garbage bytes — compression will fall back to original,
                # so the data URL still encodes the bytes we wrote in.
                zf.writestr("word/media/image10.png", b"TEN_BYTES_PAYLOAD")
                zf.writestr("word/media/image2.png", b"TWO_BYTES_PAYLOAD")
                zf.writestr("word/media/image1.png", b"ONE_BYTES_PAYLOAD")

            urls = parser.extract_docx_embedded_images_as_data_urls(path)

        # Decode the base64 portion of each data URL and check ordering.
        import base64
        prefix = "data:image/png;base64,"
        decoded = [base64.b64decode(u[len(prefix):]) for u in urls]
        self.assertEqual(decoded, [b"ONE_BYTES_PAYLOAD", b"TWO_BYTES_PAYLOAD", b"TEN_BYTES_PAYLOAD"])

    def test_compress_image_to_jpeg_shrinks_oversized_pngs(self):
        try:
            from PIL import Image
        except ImportError:
            self.skipTest("Pillow is not installed")

        import io
        buf = io.BytesIO()
        Image.new("RGB", (3000, 3000), (255, 128, 64)).save(buf, format="PNG")
        original = buf.getvalue()

        compressed, mime = parser._compress_image_to_jpeg_bytes(original)
        self.assertEqual(mime, "image/jpeg")
        self.assertLess(len(compressed), len(original))

        # And confirm the long edge was capped at 1280
        out = Image.open(io.BytesIO(compressed))
        self.assertLessEqual(max(out.size), 1280)

    def test_compress_image_falls_back_for_non_image_bytes(self):
        compressed, mime = parser._compress_image_to_jpeg_bytes(b"not an image at all")
        self.assertIsNone(mime)
        self.assertEqual(compressed, b"not an image at all")

    def test_extract_docx_embedded_images_returns_empty_for_bad_zip(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "broken.docx")
            with open(path, "wb") as f:
                f.write(b"not a zip")
            self.assertEqual(parser.extract_docx_embedded_images_as_data_urls(path), [])

    def test_parse_document_uses_text_first_then_docx_image_annotations(self):
        text_result = {"questions": [{"question_num": 4, "stem": "装置正确的是"}]}
        annotated = {"questions": [{"question_num": 4, "stem": "装置正确的是", "image_refs": [1]}]}

        with mock.patch.object(parser, "extract_text_from_word", return_value="题干文字" * 20), \
             mock.patch.object(parser, "extract_docx_embedded_images_as_data_urls",
                               return_value=["data:image/png;base64,docximg"]), \
             mock.patch.object(parser, "call_mimo", return_value=text_result) as call_mimo, \
             mock.patch.object(parser, "call_mimo_image_annotation_batches", return_value=[{"annotations": []}]) as annotate, \
             mock.patch.object(parser, "merge_image_annotations", return_value=annotated) as merge:
            result = parser.parse_document_to_questions("paper.docx")

        self.assertEqual(result["questions"], annotated["questions"])
        self.assertEqual(call_mimo.call_count, 1)
        self.assertEqual(call_mimo.call_args.args[0], "题干文字" * 20)
        annotate.assert_called_once_with("题干文字" * 20, text_result, ["data:image/png;base64,docximg"])
        merge.assert_called_once_with(text_result, [{"annotations": []}], 1)


    def test_extract_docx_embedded_images_does_not_truncate_by_default(self):
        import io
        import zipfile

        try:
            from PIL import Image
        except ImportError:
            self.skipTest("Pillow is not installed")

        def png_bytes(i):
            buf = io.BytesIO()
            Image.new("RGB", (20, 20), (i % 255, 0, 0)).save(buf, format="PNG")
            return buf.getvalue()

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "many-images.docx")
            with zipfile.ZipFile(path, "w") as zf:
                zf.writestr("word/document.xml", "<doc/>")
                for i in range(1, 31):
                    zf.writestr(f"word/media/image{i}.png", png_bytes(i))

            urls = parser.extract_docx_embedded_images_as_data_urls(path)

        self.assertEqual(len(urls), 30)
        self.assertTrue(all(u.startswith("data:image/jpeg;base64,") for u in urls))

    def test_chunk_items_splits_without_dropping_items(self):
        chunks = parser.chunk_items(list(range(25)), 12)
        self.assertEqual(chunks, [list(range(12)), list(range(12, 24)), [24]])

    def test_merge_question_batches_prefers_later_duplicate_question_numbers(self):
        batches = [
            {"questions": [
                {"question_num": 1, "stem": "text only", "answer": "A"},
                {"question_num": 2, "stem": "text q2", "answer": "B"},
            ]},
            {"questions": [
                {"question_num": 2, "stem": "image-aware q2", "answer": "C"},
                {"question_num": 3, "stem": "image-aware q3", "answer": "D"},
            ]},
        ]

        merged = parser.merge_question_batches(batches)

        self.assertEqual([q["question_num"] for q in merged["questions"]], [1, 2, 3])
        self.assertEqual(merged["questions"][1]["stem"], "image-aware q2")
        self.assertEqual(merged["questions"][1]["answer"], "C")

    def test_merge_image_annotations_adds_valid_image_refs_and_visual_notes(self):
        text_result = {"questions": [
            {"question_num": 1, "stem": "q1", "explanation": "原解析"},
            {"question_num": 2, "stem": "q2", "image_refs": [2]},
        ]}
        annotation_results = [
            {"annotations": [
                {"question_num": 1, "image_refs": [1, 1, 99, "bad"], "visual_note": "图1为实验装置"},
                {"question_num": 2, "image_refs": [2, 3], "visual_note": "图3为曲线"},
                {"question_num": 999, "image_refs": [4], "visual_note": "忽略"},
            ]}
        ]

        merged = parser.merge_image_annotations(text_result, annotation_results, image_count=3)

        self.assertEqual(merged["questions"][0]["image_refs"], [1])
        self.assertIn("原解析", merged["questions"][0]["explanation"])
        self.assertIn("图像信息：图1为实验装置", merged["questions"][0]["explanation"])
        self.assertEqual(merged["questions"][1]["image_refs"], [2, 3])
        self.assertIn("图像信息：图3为曲线", merged["questions"][1]["explanation"])

    def test_merge_image_annotations_keeps_text_result_when_annotations_error(self):
        text_result = {"questions": [{"question_num": 1, "stem": "q1"}]}

        merged = parser.merge_image_annotations(
            text_result,
            [{"error": "API call failed: timeout"}],
            image_count=2,
        )

        self.assertEqual(merged, text_result)

    def test_merge_image_annotations_sanitizes_existing_refs_without_matching_annotations(self):
        text_result = {"questions": [
            {"question_num": 1, "stem": "q1", "image_refs": [0, 2, True, 99, "bad"]},
            {"question_num": 2, "stem": "q2", "image_refs": [False, "bad"]},
        ]}

        merged = parser.merge_image_annotations(
            text_result,
            [{"annotations": [{"question_num": 999, "image_refs": [1]}]}],
            image_count=3,
        )

        self.assertEqual(merged["questions"][0]["image_refs"], [2])
        self.assertNotIn("image_refs", merged["questions"][1])

    def test_merge_image_annotations_removes_invalid_existing_refs_when_none_valid(self):
        text_result = {"questions": [
            {"question_num": 1, "stem": "q1", "image_refs": [0, 99, "bad"]},
        ]}
        annotation_results = [{"annotations": [
            {"question_num": 1, "image_refs": [0, 99, "bad"]},
        ]}]

        merged = parser.merge_image_annotations(text_result, annotation_results, image_count=3)

        self.assertNotIn("image_refs", merged["questions"][0])

    def test_merge_image_annotations_excludes_bool_refs(self):
        text_result = {"questions": [
            {"question_num": 1, "stem": "q1", "image_refs": [True]},
        ]}
        annotation_results = [{"annotations": [
            {"question_num": 1, "image_refs": [False]},
        ]}]

        merged = parser.merge_image_annotations(text_result, annotation_results, image_count=3)

        self.assertNotIn("image_refs", merged["questions"][0])

    def test_call_mimo_image_annotation_batches_sends_lightweight_annotation_prompt(self):
        image_urls = [f"data:image/jpeg;base64,{i}" for i in range(5)]
        text_result = {"questions": [
            {"question_num": 1, "stem": "第一题"},
            {"question_num": 2, "stem": "第二题"},
        ]}
        responses = [
            {"annotations": [{"question_num": 1, "image_refs": [1], "visual_note": "装置图"}]},
            {"annotations": [{"question_num": 2, "image_refs": [5], "visual_note": "曲线图"}]},
        ]

        with mock.patch.object(parser, "call_mimo", side_effect=responses) as call_mimo:
            result = parser.call_mimo_image_annotation_batches("paper text", text_result, image_urls, batch_size=4)

        self.assertEqual(result, responses)
        self.assertEqual(call_mimo.call_count, 2)
        first_prompt = call_mimo.call_args_list[0].args[0]
        self.assertIn("只输出图片标注", first_prompt)
        self.assertIn("question_num", first_prompt)
        self.assertIn("image_refs", first_prompt)
        self.assertIn("visual_note", first_prompt)
        self.assertEqual(call_mimo.call_args_list[0].kwargs["image_urls"], image_urls[:4])
        self.assertEqual(call_mimo.call_args_list[1].kwargs["image_urls"], image_urls[4:])
        self.assertEqual(call_mimo.call_args_list[0].kwargs["system_prompt"], parser.IMAGE_ANNOTATION_PROMPT)

    def test_call_mimo_image_annotation_batches_returns_error_batch_without_raising(self):
        image_urls = [f"data:image/jpeg;base64,{i}" for i in range(4)]
        text_result = {"questions": [{"question_num": 1, "stem": "第一题"}]}

        with mock.patch.object(parser, "call_mimo", return_value={"error": "API call failed: timeout"}):
            result = parser.call_mimo_image_annotation_batches("paper text", text_result, image_urls, batch_size=4)

        self.assertEqual(result, [{"error": "API call failed: timeout"}])

    def test_call_mimo_with_image_batches_logs_each_batch_timing(self):
        image_urls = [f"data:image/jpeg;base64,{i}" for i in range(9)]
        responses = [
            {"questions": [{"question_num": 1, "stem": "q1"}]},
            {"questions": [{"question_num": 2, "stem": "q2"}]},
        ]

        with self.assertLogs("parser", level="INFO") as logs, \
             mock.patch.object(parser, "call_mimo", side_effect=responses):
            parser.call_mimo_with_image_batches("paper text", image_urls, batch_size=8)

        output = "\n".join(logs.output)
        self.assertIn("mimo_batch start index=1 total=2 images=8", output)
        self.assertIn("mimo_batch done index=1 total=2 images=8", output)
        self.assertIn("mimo_batch start index=2 total=2 images=1", output)
        self.assertIn("mimo_batch done index=2 total=2 images=1", output)

    def test_call_mimo_with_image_batches_sends_all_images_in_chunks(self):
        image_urls = [f"data:image/jpeg;base64,{i}" for i in range(25)]
        responses = [
            {"questions": [{"question_num": 1, "stem": "q1"}]},
            {"questions": [{"question_num": 2, "stem": "q2"}]},
            {"questions": [{"question_num": 3, "stem": "q3"}]},
        ]

        with mock.patch.object(parser, "call_mimo", side_effect=responses) as call_mimo:
            result = parser.call_mimo_with_image_batches("paper text", image_urls, batch_size=10)

        self.assertEqual([q["question_num"] for q in result["questions"]], [1, 2, 3])
        self.assertEqual(call_mimo.call_count, 3)
        self.assertEqual(call_mimo.call_args_list[0].kwargs["image_urls"], image_urls[:10])
        self.assertEqual(call_mimo.call_args_list[1].kwargs["image_urls"], image_urls[10:20])
        self.assertEqual(call_mimo.call_args_list[2].kwargs["image_urls"], image_urls[20:])

    def test_call_mimo_with_image_batches_returns_error_when_a_batch_fails(self):
        image_urls = [f"data:image/jpeg;base64,{i}" for i in range(20)]

        with mock.patch.object(parser, "call_mimo", side_effect=[
            {"questions": [{"question_num": 1}]},
            {"error": "API HTTP 500"},
        ]):
            result = parser.call_mimo_with_image_batches("paper text", image_urls, batch_size=10)

        self.assertEqual(result, {"error": "API HTTP 500"})

    def test_call_mimo_with_image_batches_without_images_calls_call_mimo_directly(self):
        with mock.patch.object(parser, "call_mimo", return_value={"questions": [{"question_num": 1}]}) as call_mimo:
            result = parser.call_mimo_with_image_batches("paper text", image_urls=[])

        self.assertEqual(call_mimo.call_count, 1)
        args, kwargs = call_mimo.call_args
        self.assertEqual(args[0], "paper text")
        self.assertNotIn("image_urls", kwargs)

    def test_parse_document_keeps_text_questions_when_image_annotation_fails(self):
        text_result = {"questions": [{"question_num": 4, "stem": "装置正确的是"}]}
        images = ["data:image/jpeg;base64,1"]

        with mock.patch.object(parser, "extract_text_from_word", return_value="题干文字" * 20), \
             mock.patch.object(parser, "extract_docx_embedded_images_as_data_urls", return_value=images), \
             mock.patch.object(parser, "call_mimo", return_value=text_result), \
             mock.patch.object(parser, "call_mimo_image_annotation_batches", return_value=[{"error": "API call failed: timeout"}]):
            result = parser.parse_document_to_questions("paper.docx")

        self.assertEqual(result["questions"], text_result["questions"])
        self.assertEqual(result["images"], images)
        self.assertEqual(result["image_source"], "DOCX embedded images")

    def test_parse_document_logs_timing_diagnostics_for_docx_images(self):
        expected = {"questions": [{"question_num": 4, "stem": "装置正确的是"}]}
        images = [f"data:image/jpeg;base64,{i}" for i in range(9)]

        with self.assertLogs("parser", level="INFO") as logs, \
             mock.patch.object(parser, "extract_text_from_word", return_value="题干文字" * 20), \
             mock.patch.object(parser, "extract_docx_embedded_images_as_data_urls", return_value=images), \
             mock.patch.object(parser, "call_mimo", return_value=expected), \
             mock.patch.object(parser, "call_mimo_image_annotation_batches", return_value=[]):
            parser.parse_document_to_questions("paper.docx")

        output = "\n".join(logs.output)
        self.assertIn("parse_document start", output)
        self.assertIn("kind=word", output)
        self.assertIn("text_chars=80", output)
        self.assertIn("images=9", output)
        self.assertIn("batches=3", output)
        self.assertIn("parse_document done", output)

    def test_parse_document_uses_image_annotation_batching_for_many_docx_images(self):
        text_result = {"questions": [{"question_num": 4, "stem": "装置正确的是"}]}
        annotated = {"questions": [{"question_num": 4, "stem": "装置正确的是", "image_refs": [1, 2]}]}
        images = [f"data:image/jpeg;base64,{i}" for i in range(25)]

        with mock.patch.object(parser, "extract_text_from_word", return_value="题干文字" * 20), \
             mock.patch.object(parser, "extract_docx_embedded_images_as_data_urls", return_value=images), \
             mock.patch.object(parser, "call_mimo", return_value=text_result) as call_mimo, \
             mock.patch.object(parser, "call_mimo_image_annotation_batches", return_value=[{"annotations": []}]) as annotate, \
             mock.patch.object(parser, "merge_image_annotations", return_value=annotated) as merge:
            result = parser.parse_document_to_questions("paper.docx")

        self.assertEqual(result["questions"], annotated["questions"])
        self.assertEqual(call_mimo.call_count, 1)
        annotate.assert_called_once_with("题干文字" * 20, text_result, images)
        merge.assert_called_once_with(text_result, [{"annotations": []}], 25)
        self.assertEqual(result["images"], images)
        self.assertEqual(result["image_source"], "DOCX embedded images")

    def test_normalize_parse_result_filters_invalid_fields(self):
        result = parser.normalize_parse_result({
            "questions": [
                {"question_num": 2, "options": "ABCD", "image_refs": [1, 99, "x"], "stem": "题干"},
                {"question_num": None, "stem": "bad"},
            ]
        }, image_count=2)

        self.assertEqual(len(result["questions"]), 1)
        self.assertEqual(result["questions"][0]["options"], [])
        self.assertEqual(result["questions"][0]["image_refs"], [1])

    def test_question_block_parse_retries_failed_blocks_once(self):
        text = "1. 第一题题干内容足够长用于切块，并包含更多文字保证长度达标\n2. 第二题题干内容足够长用于切块，并包含更多文字保证长度达标"
        calls = []

        def fake_call(block, image_urls=None, retry=True):
            calls.append(block)
            if block.startswith("2.") and calls.count(block) == 1:
                return {"error": "temporary"}
            return {"questions": [{"question_num": int(block[0]), "stem": block}]}

        with mock.patch.object(parser, "call_mimo", side_effect=fake_call):
            result = parser.parse_question_blocks_with_mimo(text)

        self.assertEqual([q["question_num"] for q in result["questions"]], [1, 2])
        self.assertEqual(calls.count("2. 第二题题干内容足够长用于切块，并包含更多文字保证长度达标"), 2)

if __name__ == "__main__":
    unittest.main()
